"""
NexGenCV AI - Resume Parser Service
Extracts structured text, candidate profile, experience, education, projects, and skills.
"""
import re
from pypdf import PdfReader
from docx import Document
from typing import Dict, List, Any, Optional
from app.models.schemas import CandidateInfo, Project, Experience, ExperienceSummary, Education
from app.services.ocr_service import ocr_service

class ResumeParser:
    """Parse resumes and extract clean structured data from PDF & DOCX formats."""
    
    PARSING_STANDARD = "standard"
    PARSING_OCR = "ocr"
    PARSING_OCR_UNAVAILABLE = "ocr_unavailable"
    
    EMAIL_PATTERN = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    PHONE_PATTERN = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3,5}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}|\b\d{10}\b'
    LINKEDIN_PATTERN = r'(?:linkedin\.com/in/|linkedin:?\s*)([a-zA-Z0-9-_]+)'
    GITHUB_PATTERN = r'(?:github\.com/|github:?\s*)([a-zA-Z0-9-_]+)'
    
    SECTION_HEADERS = {
        'education': ['education', 'academic background', 'qualification', 'academics', 'educational background'],
        'experience': ['professional experience', 'work experience', 'experience', 'employment', 'work history', 'career history'],
        'skills': ['skills', 'technical skills', 'skills & abilities', 'competencies', 'technologies', 'tech stack', 'expertise'],
        'projects': ['projects', 'personal projects', 'key projects', 'academic projects', 'portfolio'],
        'certifications': ['certifications', 'certificates', 'credentials', 'licenses'],
        'summary': ['summary', 'profile', 'professional summary', 'objective', 'about me']
    }
    
    ACTION_VERBS = [
        'achieved', 'administered', 'analyzed', 'architected', 'automated',
        'built', 'collaborated', 'configured', 'created', 'delivered',
        'designed', 'developed', 'drove', 'enhanced', 'established',
        'executed', 'implemented', 'improved', 'increased', 'integrated',
        'launched', 'led', 'managed', 'mentored', 'migrated',
        'optimized', 'orchestrated', 'oversaw', 'pioneered', 'planned',
        'reduced', 'refactored', 'resolved', 'scaled', 'secured',
        'spearheaded', 'streamlined', 'supervised', 'transformed', 'upgraded'
    ]

    def parse(self, file_path: str, file_ext: str) -> Dict[str, Any]:
        parsing_method = self.PARSING_STANDARD
        ocr_confidence = None
        
        if file_ext == '.pdf':
            raw_text = self._extract_pdf_text(file_path)
            has_tables = self._check_pdf_tables(file_path)
            has_images = self._check_pdf_images(file_path)
            raw_text, parsing_method, ocr_confidence = self._apply_ocr_if_needed(file_path, raw_text)
        else:
            raw_text = self._extract_docx_text(file_path)
            has_tables = self._check_docx_tables(file_path)
            has_images = self._check_docx_images(file_path)
        
        sections = self._identify_sections(raw_text)
        candidate = self._extract_candidate_info(raw_text)
        experience = self._extract_experience(raw_text, sections.get('experience', ''))
        projects = self._extract_projects(raw_text, sections.get('projects', ''))
        education = self._extract_education(raw_text, sections.get('education', ''))
        
        return {
            "raw_text": raw_text,
            "candidate": candidate,
            "experience": experience,
            "projects": projects,
            "education": education,
            "sections": sections,
            "formatting": {
                "has_tables": has_tables,
                "has_images": has_images,
                "word_count": len(raw_text.split()),
                "line_count": len(raw_text.split('\n'))
            },
            "parsing_method": parsing_method,
            "ocr_confidence": ocr_confidence
        }

    def _apply_ocr_if_needed(self, file_path: str, standard_text: str) -> tuple:
        if not ocr_service.is_available():
            return standard_text, self.PARSING_STANDARD, None
        
        email_match = re.search(self.EMAIL_PATTERN, standard_text)
        phone_match = re.search(self.PHONE_PATTERN, standard_text)
        
        if not ocr_service.needs_ocr(
            standard_text, 
            email=email_match.group() if email_match else None,
            phone=phone_match.group() if phone_match else None
        ):
            return standard_text, self.PARSING_STANDARD, None
        
        if ocr_service.should_skip_ocr(file_path):
            return standard_text, self.PARSING_OCR_UNAVAILABLE, None
        
        ocr_text, parsing_method, confidence = ocr_service.extract_text_with_ocr(file_path)
        if ocr_text and parsing_method == self.PARSING_OCR:
            return ocr_text, parsing_method, confidence
        return standard_text, parsing_method, confidence

    def _extract_pdf_text(self, file_path: str) -> str:
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
        return text

    def _extract_docx_text(self, file_path: str) -> str:
        text = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
        return text

    def _check_pdf_tables(self, file_path: str) -> bool:
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                text = page.extract_text() or ""
                lines = text.split('\n')
                table_like_lines = sum(1 for line in lines if line.count('\t') >= 2 or line.count('|') >= 2)
                if table_like_lines > 3:
                    return True
        except:
            pass
        return False

    def _check_pdf_images(self, file_path: str) -> bool:
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                if '/XObject' in page.get('/Resources', {}):
                    xobject = page['/Resources']['/XObject']
                    if xobject:
                        for obj in xobject:
                            if xobject[obj]['/Subtype'] == '/Image':
                                return True
        except:
            pass
        return False

    def _check_docx_tables(self, file_path: str) -> bool:
        try:
            doc = Document(file_path)
            return len(doc.tables) > 0
        except:
            return False

    def _check_docx_images(self, file_path: str) -> bool:
        try:
            doc = Document(file_path)
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    return True
        except:
            return False

    def _identify_sections(self, text: str) -> Dict[str, str]:
        sections = {}
        lines = text.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            line_clean = line.strip()
            line_lower = line_clean.lower()
            section_found = None
            
            # Match section headers
            for section_type, headers in self.SECTION_HEADERS.items():
                for header in headers:
                    if line_lower == header or line_lower == header + ':' or line_lower.startswith(header + ' ') and len(line_lower) < 35:
                        section_found = section_type
                        break
                if section_found:
                    break
            
            if section_found:
                if current_section:
                    sections[current_section] = '\n'.join(current_content)
                current_section = section_found
                current_content = []
            elif current_section:
                current_content.append(line)
        
        if current_section:
            sections[current_section] = '\n'.join(current_content)
        
        return sections

    def _extract_candidate_info(self, text: str) -> CandidateInfo:
        lines = [l.strip() for l in text.split('\n') if l.strip()][:10]
        name = None
        
        # Candidate Name Detection (usually Line 1)
        for line in lines:
            if '@' in line or re.search(self.PHONE_PATTERN, line):
                continue
            if any(kw in line.lower() for kw in ['resume', 'curriculum', 'education', 'skills', 'experience']):
                continue
            words = line.split()
            if 1 <= len(words) <= 4 and all(w.replace('.', '').replace('-', '').isalpha() for w in words):
                name = line
                break
        
        email_match = re.search(self.EMAIL_PATTERN, text)
        email = email_match.group() if email_match else None
        
        phone_match = re.search(self.PHONE_PATTERN, text)
        phone = phone_match.group() if phone_match else None
        
        linkedin_match = re.search(self.LINKEDIN_PATTERN, text, re.IGNORECASE)
        linkedin = f"linkedin.com/in/{linkedin_match.group(1)}" if linkedin_match else None
        
        github_match = re.search(self.GITHUB_PATTERN, text, re.IGNORECASE)
        github = f"github.com/{github_match.group(1)}" if github_match else None
        
        # Check for inline GitHub handle in top lines if no full URL
        if not github:
            for line in lines[:10]:
                gh_m = re.search(r'(?:github[:\s]+|github\.com/)([a-zA-Z0-9-_]+)', line, re.IGNORECASE)
                if gh_m and 'gmail' not in gh_m.group(1).lower() and 'http' not in gh_m.group(1).lower():
                    github = f"github.com/{gh_m.group(1)}"
                    break

        location = self._extract_location(text)
        
        return CandidateInfo(
            name=name,
            email=email,
            phone=phone,
            location=location,
            linkedin=linkedin,
            github=github
        )

    def _extract_location(self, text: str) -> Optional[str]:
        known_places = {
            'india', 'usa', 'us', 'uk', 'canada', 'germany', 'australia', 'singapore', 'uae',
            'gujarat', 'maharashtra', 'delhi', 'karnataka', 'tamil nadu', 'telangana', 'uttar pradesh',
            'gandhinagar', 'ahmedabad', 'mumbai', 'pune', 'bangalore', 'bengaluru', 'hyderabad',
            'chennai', 'noida', 'gurugram', 'gurgaon', 'kolkata', 'surat', 'vadodara', 'jaipur',
            'new york', 'san francisco', 'london', 'toronto', 'berlin', 'sydney', 'chicago', 'seattle'
        }

        top_lines = [l.strip() for l in text[:1200].split('\n') if l.strip()][:15]

        # 1. Explicit word-boundary label match: "Location: Gandhinagar", "Address: Mumbai", "Based in: Delhi"
        for line in top_lines:
            label_match = re.search(r'\b(?:location|address|based in|city)\b[:\s]+([A-Za-z\s,.-]+)', line, re.IGNORECASE)
            if label_match:
                candidate = label_match.group(1).strip()
                candidate = re.sub(r'^[^\w]+|[^\w]+$', '', candidate)
                if 3 <= len(candidate) <= 40 and not self._is_programming_language_line(candidate):
                    return candidate

        # 2. Strict City, Country / State pattern matching against known_places list
        for line in top_lines:
            if self._is_programming_language_line(line):
                continue
            
            match = re.search(r'\b([A-Z][a-zA-Z\s]{1,20},\s*([A-Za-z\s]{2,20}))\b', line)
            if match:
                full_loc = match.group(1).strip()
                target_part = match.group(2).strip().lower()
                first_part = match.group(1).split(',')[0].strip().lower()

                is_valid = (
                    target_part in known_places or
                    first_part in known_places or
                    (len(target_part) == 2 and target_part.isupper())
                )

                if is_valid and not self._is_programming_language_line(full_loc):
                    return full_loc

        return None

    def _is_programming_language_line(self, text: str) -> bool:
        text_lower = text.lower()
        tech_words = [
            'javascript', 'python', 'java', 'c++', 'c#', 'html', 'css', 'react', 'node',
            'sql', 'express', 'git', 'mongodb', 'typescript', 'php', 'ruby', 'programming',
            'languages', 'skills', 'frameworks', 'libraries', 'tools', 'databases', 'c'
        ]
        words = re.findall(r'\b[a-z+#]+\b', text_lower)
        return any(w in tech_words for w in words)

    def _extract_education(self, full_text: str, education_section: str) -> List[Education]:
        education_list = []
        text_to_analyze = education_section if education_section else full_text
        lines = [l.strip() for l in text_to_analyze.split('\n') if l.strip()]
        
        edu_keywords = ['b.e', 'b.tech', 'b.s', 'b.a', 'm.tech', 'm.s', 'hsc', 'ssc', 'bachelor', 'master', 'diploma', 'degree', 'school', 'college', 'university', 'institute']
        
        current_item = None
        
        for line in lines:
            line_lower = line.lower()
            is_edu_line = any(kw in line_lower for kw in edu_keywords)
            
            if is_edu_line:
                if current_item and current_item.get('degree'):
                    education_list.append(Education(**current_item))
                
                parts = [p.strip() for p in line.split(',') if p.strip()]
                degree = parts[0] if parts else line
                institution = parts[1] if len(parts) > 1 else None
                
                # Clean dates from institution
                if institution:
                    institution = re.sub(r'\b(20\d{2}\s*[-–]\s*20\d{2}|20\d{2}|19\d{2})\b.*', '', institution).strip(' |')
                
                year_match = re.search(r'\b(20\d{2}\s*[-–]\s*20\d{2}|20\d{2}|19\d{2})\b', line)
                year = year_match.group() if year_match else None
                
                cgpa_match = re.search(r'(?:CGPA|GPA|PCM|Percentage|Marks)[:\s–-]*([0-9.]+%?)', line, re.IGNORECASE)
                gpa = cgpa_match.group(1) if cgpa_match else None
                
                current_item = {
                    'degree': degree,
                    'institution': institution,
                    'year': year,
                    'gpa': gpa
                }
            elif current_item:
                cgpa_match = re.search(r'(?:CGPA|GPA|PCM|Percentage|Marks)[:\s–-]*([0-9.]+%?)', line, re.IGNORECASE)
                if cgpa_match and not current_item.get('gpa'):
                    current_item['gpa'] = cgpa_match.group(1)
                elif not current_item.get('institution') and len(line) < 80:
                    current_item['institution'] = line
                elif not current_item.get('year'):
                    year_match = re.search(r'\b(20\d{2}\s*[-–]\s*20\d{2}|20\d{2})\b', line)
                    if year_match:
                        current_item['year'] = year_match.group()

        if current_item and current_item.get('degree'):
            education_list.append(Education(**current_item))
            
        return education_list[:4]

    def _extract_experience(self, full_text: str, experience_section: str) -> ExperienceSummary:
        positions = []
        text_to_analyze = experience_section if experience_section else full_text
        entries = self._split_experience_entries(text_to_analyze)
        
        total_months = 0
        for entry in entries[:5]:
            exp = self._parse_experience_entry(entry)
            if exp.company or exp.role:
                positions.append(exp)
                if exp.duration:
                    total_months += self._estimate_duration_months(exp.duration)
        
        overall_quality = 75
        if positions:
            overall_quality = sum(p.bullet_quality for p in positions) // len(positions)
        
        return ExperienceSummary(
            total_years=round(max(0.5, total_months / 12), 1),
            total_months=total_months,
            positions=positions,
            overall_quality=overall_quality
        )

    def _split_experience_entries(self, text: str) -> List[str]:
        lines = [l for l in text.split('\n') if l.strip()]
        entries = []
        current_entry = []
        
        for line in lines:
            # Check if line starts a new job entry (Role – Company or Date)
            date_match = re.search(r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}', line, re.IGNORECASE)
            is_header = ('–' in line or '-' in line or '|' in line) and not line.strip().startswith(('•', '-', '*'))
            
            if (date_match or is_header) and current_entry and len(current_entry) >= 2:
                entries.append('\n'.join(current_entry))
                current_entry = []
            current_entry.append(line)
            
        if current_entry:
            entries.append('\n'.join(current_entry))
        return entries if entries else [text]

    def _parse_experience_entry(self, entry: str) -> Experience:
        lines = [l.strip() for l in entry.split('\n') if l.strip()]
        
        company = None
        role = None
        duration = None
        bullets = []
        
        for line in lines:
            if line.startswith(('•', '-', '*', '●')):
                bullets.append(line)
                continue
                
            # Date line
            date_match = re.search(
                r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4})\s*[-–—to]+\s*(?:Present|Current|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4})',
                line, re.IGNORECASE
            )
            if date_match:
                duration = date_match.group()
            
            # Header line: Role – Company
            if ('–' in line or '-' in line or '|' in line) and not role:
                parts = re.split(r'[-–|]', line)
                if len(parts) >= 2:
                    role = parts[0].strip()
                    # Clean company name from dates if present
                    comp_part = parts[1].strip()
                    comp_clean = re.sub(r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}.*', '', comp_part, flags=re.IGNORECASE).strip()
                    company = comp_clean if comp_clean else comp_part
            elif not role and any(kw in line.lower() for kw in ['developer', 'engineer', 'intern', 'manager', 'lead', 'analyst', 'designer']):
                role = line
            elif not company and len(line) < 50 and not date_match:
                company = line

        action_count = 0
        has_metrics = False
        for b in bullets:
            b_lower = b.lower()
            if any(v in b_lower for v in self.ACTION_VERBS):
                action_count += 1
            if re.search(r'\d+%|\$\d+|improved|increased|reduced', b_lower):
                has_metrics = True
                
        quality = 60 + (action_count * 10) + (20 if has_metrics else 0)
        
        return Experience(
            company=company or "Software Industry",
            role=role or "Software Engineer / Developer",
            duration=duration or "2025 – Present",
            description='\n'.join(bullets[:6]) if bullets else entry,
            bullet_quality=min(100, quality),
            has_metrics=has_metrics,
            action_verbs_count=action_count
        )

    def _estimate_duration_months(self, duration_str: str) -> int:
        if 'present' in duration_str.lower() or 'current' in duration_str.lower():
            return 6
        dates = re.findall(r'\b(20\d{2}|19\d{2})\b', duration_str)
        if len(dates) >= 2:
            try:
                y1, y2 = int(dates[0]), int(dates[1])
                return max(1, (y2 - y1) * 12)
            except:
                pass
        return 3

    def _extract_projects(self, full_text: str, projects_section: str) -> List[Project]:
        projects = []
        text_to_analyze = projects_section if projects_section else full_text
        entries = self._split_project_entries(text_to_analyze)
        
        for entry in entries[:6]:
            proj = self._parse_project_entry(entry)
            if proj.title:
                projects.append(proj)
        
        return projects

    def _split_project_entries(self, text: str) -> List[str]:
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        entries = []
        current_entry = []
        
        for line in lines:
            # Match project title lines e.g. Project Name — Subtitle or Tech Stack:
            is_title = ('—' in line or '-' in line or 'Tech Stack:' in line or 'Tech:' in line) and not line.startswith(('•', '-', '*'))
            if is_title and current_entry and len(current_entry) >= 2:
                entries.append('\n'.join(current_entry))
                current_entry = []
            current_entry.append(line)
            
        if current_entry:
            entries.append('\n'.join(current_entry))
        return entries

    def _parse_project_entry(self, entry: str) -> Project:
        lines = [l.strip() for l in entry.split('\n') if l.strip()]
        
        title = None
        technologies = []
        description_bullets = []
        
        for i, line in enumerate(lines):
            # Check if line is a tech stack declaration (regardless of position)
            is_tech_line = bool(re.search(r'^\s*(?:Tech\s+Stack|Tech|Stack|Technologies)[:\s]+', line, re.IGNORECASE))
            
            if is_tech_line:
                tech_str = re.sub(r'^(?:Tech\s+Stack|Tech|Stack|Technologies)[:\s]+', '', line, flags=re.IGNORECASE).strip()
                techs = [t.strip() for t in tech_str.split(',') if t.strip()]
                technologies.extend(techs)
            elif i == 0 and not title:
                # First non-tech-stack line is the title
                title = line.replace('•', '').replace('-', '').strip()
            elif line.startswith(('•', '-', '*', '●')):
                description_bullets.append(line)
            elif not technologies and any(tech in line.lower() for tech in ['react', 'node', 'express', 'mongodb', 'python', 'html', 'css', 'tailwind', 'jwt', 'sql', 'javascript']):
                # Implicit tech list line (no prefix)
                techs = [t.strip() for t in line.split(',') if t.strip()]
                technologies.extend(techs)
            elif title:
                description_bullets.append(line)
                
        score = 60 + (min(len(technologies), 5) * 5) + (min(len(description_bullets), 5) * 5)
        
        return Project(
            title=title or "Software Project",
            technologies=list(dict.fromkeys(technologies)),  # preserve order, deduplicate
            description='\n'.join(description_bullets[:5]) if description_bullets else None,
            impact="Full-stack implementation with clean architecture & user workflow optimization.",
            score=min(100, score)
        )
